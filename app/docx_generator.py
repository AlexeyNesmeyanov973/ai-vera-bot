# app/docx_generator.py
import logging
import os
import hashlib
from typing import List, Dict, Iterable, Optional, Tuple

logger = logging.getLogger(__name__)

# Пытаемся импортировать python-docx; при неудаче — мягко отключаем функциональность
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except Exception:
    Document = None  # type: ignore
    Pt = RGBColor = WD_STYLE_TYPE = WD_ALIGN_PARAGRAPH = None  # type: ignore
    _DOCX_AVAILABLE = False
    logger.warning("python-docx не установлен. Добавьте 'python-docx' в requirements.txt")


# ---------------- ВСПОМОГАТЕЛЬНЫЕ ----------------

def _fmt_hms(v: Optional[float]) -> str:
    if v is None:
        return "--:--:--"
    try:
        total = max(0, int(round(float(v))))
        h = total // 3600
        m = (total % 3600) // 60
        s = total % 60
        return f"{h:02}:{m:02}:{s:02}"
    except Exception:
        return "--:--:--"

def _ensure_parent_dir(path: str) -> None:
    d = os.path.dirname(os.path.abspath(path))
    if d:
        os.makedirs(d, exist_ok=True)

def _norm_text(s: str) -> str:
    return " ".join((s or "").replace("\u200b", "").split()).strip()

def _split_to_paragraphs(text: str) -> Iterable[str]:
    t = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    parts = [p.strip() for p in t.split("\n") if p.strip()]
    return parts or [""]


def _palette() -> List["RGBColor"]:  # type: ignore[name-defined]
    """
    Фиксированная приятная палитра. Создаём только если docx доступен.
    """
    if not _DOCX_AVAILABLE:  # защитный рантайм-гард
        return []
    return [
        RGBColor(33, 150, 243),   # синий
        RGBColor(244, 67, 54),    # красный
        RGBColor(76, 175, 80),    # зелёный
        RGBColor(255, 193, 7),    # янтарный
        RGBColor(156, 39, 176),   # фиолетовый
        RGBColor(0, 188, 212),    # бирюзовый
        RGBColor(255, 87, 34),    # оранжевый
        RGBColor(121, 85, 72),    # коричневый
        RGBColor(63, 81, 181),    # индиго
        RGBColor(139, 195, 74),   # лайм
        RGBColor(0, 150, 136),    # teal
        RGBColor(233, 30, 99),    # розовый
    ]

def _stable_color_for_speaker(speaker: str) -> "RGBColor":  # type: ignore[name-defined]
    """
    Детерминированное сопоставление «спикер → цвет» (md5, без прыжков между рестартами).
    """
    pal = _palette()
    if not pal:
        # вернём первый-лучший цвет, но сюда попадём только если _DOCX_AVAILABLE=False,
        # а генерация в любом случае вернёт False раньше.
        return None  # type: ignore[return-value]
    h = hashlib.md5((speaker or "SPK").encode("utf-8")).hexdigest()
    idx = int(h[:8], 16) % len(pal)
    return pal[idx]

def _speaker_key(seg: Dict) -> str:
    spk = (seg.get("speaker") or "").strip()
    return spk if spk else "SPK"


def _group_contiguous_by_speaker(segments: List[Dict]) -> List[Dict]:
    """
    Склеивает подряд идущие сегменты одного спикера в группы:
    [{speaker, start, end, texts: [..]}]
    """
    groups: List[Dict] = []
    cur: Optional[Dict] = None

    for seg in segments or []:
        text = _norm_text(seg.get("text") or "")
        if not text:
            continue
        spk = _speaker_key(seg)

        st = seg.get("start")
        en = seg.get("end")
        try:
            start = float(st) if st is not None else None
        except Exception:
            start = None
        try:
            end = float(en) if en is not None else None
        except Exception:
            end = None

        if cur and cur["speaker"] == spk:
            # расширяем блок
            if start is not None:
                cur["start"] = start if cur["start"] is None else min(cur["start"], start)
            if end is not None:
                cur["end"] = end if cur["end"] is None else max(cur["end"], end)
            cur["texts"].append(text)
        else:
            if cur:
                groups.append(cur)
            cur = {"speaker": spk, "start": start, "end": end, "texts": [text]}

    if cur:
        groups.append(cur)
    return groups

def _collect_unique_speakers_in_order(segments: List[Dict]) -> List[str]:
    order: List[str] = []
    seen = set()
    for seg in segments or []:
        spk = _speaker_key(seg)
        if spk not in seen:
            seen.add(spk)
            order.append(spk)
    return order


# ---------------- ОФОРМЛЕНИЕ WORD ----------------

def _ensure_styles(doc, base_font: str = "DejaVu Sans") -> None:
    """
    Настраиваем базовые стили. Надёжно обрабатываем отсутствие предопределённых.
    """
    styles = doc.styles

    # Normal
    normal = styles["Normal"]
    normal.font.name = base_font
    normal.font.size = Pt(11)

    # Title (best-effort)
    try:
        styles["Title"].font.name = base_font
    except KeyError:
        pass

    # SpeakerHeading
    try:
        sp = styles.add_style("SpeakerHeading", WD_STYLE_TYPE.PARAGRAPH)
        try:
            sp.base_style = styles["Heading 2"]
        except KeyError:
            sp.base_style = styles["Normal"]
        sp.font.name = base_font
        sp.font.bold = True
        sp.font.size = Pt(13)
    except ValueError:
        # уже есть
        sp = styles["SpeakerHeading"]
        sp.font.name = base_font

    # LegendHeading
    try:
        lh = styles.add_style("LegendHeading", WD_STYLE_TYPE.PARAGRAPH)
        try:
            lh.base_style = styles["Heading 3"]
        except KeyError:
            lh.base_style = styles["Normal"]
        lh.font.name = base_font
        lh.font.bold = True
        lh.font.size = Pt(12)
    except ValueError:
        lh = styles["LegendHeading"]
        lh.font.name = base_font

    # LegendEntry
    try:
        le = styles.add_style("LegendEntry", WD_STYLE_TYPE.PARAGRAPH)
        le.base_style = styles["Normal"]
        le.font.name = base_font
        le.font.size = Pt(11)
    except ValueError:
        le = styles["LegendEntry"]
        le.font.name = base_font


def _add_title(doc, title: str) -> None:
    h = doc.add_heading(title or "Транскрибация", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("")  # небольшой отступ

def _add_legend(doc, speakers: List[str], marker_char: str = "●") -> None:
    if not speakers:
        return
    doc.add_paragraph("Легенда", style="LegendHeading")
    for spk in speakers:
        line = doc.add_paragraph(style="LegendEntry")
        run_marker = line.add_run(f"{marker_char} ")
        color = _stable_color_for_speaker(spk)
        if color is not None:
            run_marker.font.color.rgb = color
        run_marker.bold = True

        run_name = line.add_run(spk)
        run_name.bold = True
    doc.add_paragraph("")

def _write_group(doc, group: Dict, with_timestamps: bool, marker_char: str) -> None:
    """
    Один блок: подзаголовок-спикер + (опц.) таймкоды + абзацы речи.
    """
    speaker: str = group["speaker"]
    start: Optional[float] = group["start"]
    end: Optional[float] = group["end"]
    texts: List[str] = group["texts"]

    # Заголовок
    head = doc.add_paragraph(style="SpeakerHeading")
    color = _stable_color_for_speaker(speaker)

    r_marker = head.add_run(f"{marker_char} ")
    if color is not None:
        r_marker.font.color.rgb = color
    r_marker.bold = True

    r_name = head.add_run(speaker)
    r_name.bold = True

    if with_timestamps and (start is not None or end is not None):
        head.add_run(f"  [{_fmt_hms(start)}–{_fmt_hms(end)}]").italic = True

    # Текст (делим на абзацы)
    body_text = " ".join(texts).strip()
    for para in _split_to_paragraphs(body_text):
        p = doc.add_paragraph()
        p.add_run(para)

    doc.add_paragraph("")  # отступ между блоками


# ---------------- ПУБЛИЧНЫЙ ИНТЕРФЕЙС ----------------

class DOCXGenerator:
    """
    Генератор DOCX:
      - generate_plain_docx(text, output_path, title)
      - generate_speaker_docx(segments, output_path, title, with_timestamps, show_legend, marker_char)
    """

    def generate_plain_docx(self, text: str, output_path: str, title: str = "Транскрибация") -> bool:
        if not _DOCX_AVAILABLE:
            logger.error("python-docx не установлен")
            return False
        try:
            _ensure_parent_dir(output_path)
            doc = Document()
            _ensure_styles(doc)

            _add_title(doc, title)

            # Абзацы по переносам
            for para in _split_to_paragraphs(text or ""):
                doc.add_paragraph(para)

            doc.save(output_path)
            return True
        except Exception:
            logger.exception("Ошибка генерации DOCX (plain)")
            return False

    def generate_speaker_docx(
        self,
        segments: List[Dict],
        output_path: str,
        title: str = "Транскрибация",
        with_timestamps: bool = True,
        show_legend: bool = True,
        marker_char: str = "●",
    ) -> bool:
        if not _DOCX_AVAILABLE:
            logger.error("python-docx не установлен")
            return False
        try:
            # Если нет диаризации — делаем обычный DOCX
            if not segments or not any((seg.get("speaker") or "").strip() for seg in segments):
                full_text = " ".join(_norm_text(s.get("text") or "") for s in segments or []).strip()
                return self.generate_plain_docx(full_text, output_path, title=title)

            _ensure_parent_dir(output_path)
            doc = Document()
            _ensure_styles(doc)

            _add_title(doc, title)

            if show_legend:
                speakers = _collect_unique_speakers_in_order(segments)
                _add_legend(doc, speakers, marker_char=marker_char)

            groups = _group_contiguous_by_speaker(segments)
            for g in groups:
                _write_group(doc, g, with_timestamps=with_timestamps, marker_char=marker_char)

            doc.save(output_path)
            return True
        except Exception:
            logger.exception("Ошибка генерации DOCX (speakers)")
            return False


# Инстанс и обёртки (совместимо с вашим импортом: from app.docx_generator import docx_generator)
docx_generator = DOCXGenerator()

def generate_plain_docx(text: str, output_path: str, title: str = "Транскрибация") -> bool:
    return docx_generator.generate_plain_docx(text, output_path, title=title)

def generate_speaker_docx(
    segments: List[Dict],
    output_path: str,
    title: str = "Транскрибация",
    with_timestamps: bool = True,
    show_legend: bool = True,
    marker_char: str = "●",
) -> bool:
    return docx_generator.generate_speaker_docx(
        segments=segments,
        output_path=output_path,
        title=title,
        with_timestamps=with_timestamps,
        show_legend=show_legend,
        marker_char=marker_char,
    )
